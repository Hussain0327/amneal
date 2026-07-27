"""API contract for durable White-Paper runs (design doc Phase 2, section 5).

Covers the org-shared run surface end to end over the real store + tmp SQLite:
list/detail, attributed cell edits, finalize/reopen freeze semantics, the
server-side docx render (fingerprint re-verified from STORAGE, overlay applied),
creator-only delete, and the stored-corruption 500 paths that replaced the
legacy client-echo integrity check.
"""

from __future__ import annotations

import copy
from io import BytesIO
from typing import Any, NoReturn

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlmodel import select

import regwatch.api.main as api_main
from regwatch.api.main import app
from regwatch.store.db import session_scope
from regwatch.store.models import QueryLog, WhitepaperRun
from tests._whitepaper_stub import APPL_NO, RLD_NAME, install_fake_sources
from tests.conftest import create_user, session_client

ANALYST_CELL = "rd_center"  # manual cell: analyst_input_required in the stub build
POPULATED_CELL = "product_name"  # auto cell: populated in the stub build

OTHER_EMAIL = "other@example.com"
OTHER_PASSWORD = "other-password-123"
OTHER_NAME = "Second Analyst"


def _populate(client: TestClient) -> dict[str, Any]:
    r = client.post("/whitepaper", json={"rld_name": RLD_NAME, "application_number": APPL_NO})
    assert r.status_code == 200, r.text
    body = r.json()
    assert isinstance(body["run_id"], int)
    return body


def _whitepaper_audit_rows() -> list[dict[str, Any]]:
    with session_scope() as s:
        rows = s.scalars(select(QueryLog).where(QueryLog.mode == "whitepaper"))
        return [
            {
                "status": r.status,
                "model_name": r.model_name,
                "query_text": r.query_text,
                "route_json": dict(r.route_json),
            }
            for r in rows
        ]


def _boom(*args: Any, **kwargs: Any) -> NoReturn:
    """Simulated DB outage on an audit-side round-trip (tests/test_grounded_qa_citations
    uses the same idiom for the QA path)."""
    raise RuntimeError("simulated audit db outage")


def _forbid_repopulate(monkeypatch: pytest.MonkeyPatch) -> None:
    """The saved-run endpoints must never re-populate: no fetches, no LLM."""

    def boom(*args: Any, **kwargs: Any) -> NoReturn:
        raise AssertionError("saved-run endpoint re-populated (live fetch/LLM path was invoked)")

    monkeypatch.setattr(api_main, "build_whitepaper", boom)


def _use_absent_template(monkeypatch: pytest.MonkeyPatch, tmp_path: Any) -> None:
    """Deterministic docx path: the real gitignored template may exist on a dev
    machine; force the from-scratch fallback and no fetch URL."""
    import config.settings as cs

    monkeypatch.setenv("WHITEPAPER_TEMPLATE_PATH", str(tmp_path / "absent-template.docx"))
    monkeypatch.delenv("WHITEPAPER_TEMPLATE_URL", raising=False)
    cs.get_settings.cache_clear()


def _other_client() -> TestClient:
    return session_client(create_user(OTHER_EMAIL, OTHER_PASSWORD, display_name=OTHER_NAME))


# ---------------------------------------------------------------------------
# auth (router-level) + 404s
# ---------------------------------------------------------------------------


def test_run_endpoints_require_auth() -> None:
    client = TestClient(app)
    client.__enter__()
    try:
        assert client.get("/whitepaper/runs").status_code == 401
        assert client.get("/whitepaper/runs/1").status_code == 401
        assert (
            client.post("/whitepaper/runs/1/cells/rd_center", json={"value": "x"}).status_code
            == 401
        )
        assert client.post("/whitepaper/runs/1/finalize").status_code == 401
        assert client.post("/whitepaper/runs/1/reopen").status_code == 401
        assert client.post("/whitepaper/runs/1/docx").status_code == 401
        assert client.delete("/whitepaper/runs/1").status_code == 401
    finally:
        client.__exit__(None, None, None)


def test_missing_run_404_on_every_route(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    assert auth_client.get("/whitepaper/runs/99999").status_code == 404
    r = auth_client.post("/whitepaper/runs/99999/cells/rd_center", json={"value": "x"})
    assert r.status_code == 404
    assert auth_client.post("/whitepaper/runs/99999/finalize").status_code == 404
    assert auth_client.post("/whitepaper/runs/99999/reopen").status_code == 404
    assert auth_client.post("/whitepaper/runs/99999/docx").status_code == 404
    assert auth_client.delete("/whitepaper/runs/99999").status_code == 404


# ---------------------------------------------------------------------------
# list + detail (org-shared)
# ---------------------------------------------------------------------------


def test_list_runs_shape_filters_and_pagination(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    first = _populate(auth_client)
    second = _populate(auth_client)

    r = auth_client.get("/whitepaper/runs")
    assert r.status_code == 200, r.text
    body = r.json()
    assert {"count", "total", "limit", "offset", "runs"} <= set(body)
    assert body["count"] == body["total"] == 2
    assert body["limit"] == 50 and body["offset"] == 0
    newest = body["runs"][0]
    # Newest activity first; summaries carry counts + attribution, never payloads.
    assert newest["id"] == second["run_id"]
    assert newest["application_number"] == APPL_NO
    assert newest["status"] == "draft"
    assert newest["created_by"] == "Test Analyst"
    assert newest["populated_count"] > 0
    assert newest["inputs_count"] == 0
    assert "sections" not in newest

    # Filters accept the prefixed form (normalize_appl_no) and pagination is truthful.
    filtered = auth_client.get("/whitepaper/runs", params={"application_number": f"NDA {APPL_NO}"})
    assert filtered.json()["total"] == 2
    page = auth_client.get("/whitepaper/runs", params={"limit": 1, "offset": 1}).json()
    assert page["count"] == 1 and page["total"] == 2
    assert page["runs"][0]["id"] == first["run_id"]
    none = auth_client.get("/whitepaper/runs", params={"normalized_name": "ibuprofen"}).json()
    assert none["total"] == 0
    assert auth_client.get("/whitepaper/runs", params={"status": "final"}).json()["total"] == 0


def test_list_runs_unparseable_application_number_422(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    r = auth_client.get("/whitepaper/runs", params={"application_number": "no digits here"})
    assert r.status_code == 422


def test_get_run_detail_is_verbatim_storage(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The detail response ships the generated payload VERBATIM (INV-3): the
    fingerprinted sections are passthrough, never reshaped by serialization."""
    install_fake_sources(monkeypatch)
    body = _populate(auth_client)
    r = auth_client.get(f"/whitepaper/runs/{body['run_id']}")
    assert r.status_code == 200, r.text
    detail = r.json()
    assert detail["sections"] == body["sections"]
    assert detail["spine"] == body["spine"]
    assert detail["warnings"] == body["warnings"]
    assert detail["source_audit_id"] == body["audit_id"]
    assert detail["status"] == "draft"
    assert detail["created_by"] == "Test Analyst"
    assert detail["finalized_at"] is None and detail["finalized_by"] is None
    assert detail["inputs"] == {}
    counts = (
        detail["populated_count"] + detail["analyst_input_count"] + detail["verified_absent_count"]
    )
    assert counts == sum(len(s["cells"]) for s in body["sections"])


def test_org_shared_read_and_attributed_edit(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """User B lists and reads user A's run and edits an analyst cell; the value
    is attributed to B (org-shared workflow, design doc section 10)."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    other = _other_client()
    try:
        listing = other.get("/whitepaper/runs")
        assert listing.json()["total"] == 1
        assert listing.json()["runs"][0]["created_by"] == "Test Analyst"
        assert other.get(f"/whitepaper/runs/{run_id}").status_code == 200
        r = other.post(
            f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "Filled by B"}
        )
        assert r.status_code == 200, r.text
        saved = r.json()
        assert saved["cleared"] is False
        assert saved["input"]["value"] == "Filled by B"
        assert saved["input"]["author"] == OTHER_NAME
    finally:
        other.__exit__(None, None, None)
    # A sees B's attributed input on the shared run.
    detail = auth_client.get(f"/whitepaper/runs/{run_id}").json()
    assert detail["inputs"][ANALYST_CELL]["value"] == "Filled by B"
    assert detail["inputs"][ANALYST_CELL]["author"] == OTHER_NAME
    # inputs_count rides into the summary row.
    assert auth_client.get("/whitepaper/runs").json()["runs"][0]["inputs_count"] == 1


# ---------------------------------------------------------------------------
# cell edits: clear semantics + 422s
# ---------------------------------------------------------------------------


def test_cell_upsert_clear_and_validation(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]

    r = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "Site A"}
    )
    assert r.status_code == 200 and r.json()["input"]["value"] == "Site A"

    # null clears (deletes the overlay row); a whitespace-only value is a clear too.
    cleared = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": None}
    )
    assert cleared.status_code == 200
    assert cleared.json() == {
        "run_id": run_id,
        "cell_id": ANALYST_CELL,
        "cleared": True,
        "input": None,
    }
    assert auth_client.get(f"/whitepaper/runs/{run_id}").json()["inputs"] == {}
    blank = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "   "}
    )
    assert blank.status_code == 200 and blank.json()["cleared"] is True

    # 422: unknown cell id (not in template.CELL_SPECS).
    bad_cell = auth_client.post(f"/whitepaper/runs/{run_id}/cells/not_a_cell", json={"value": "x"})
    assert bad_cell.status_code == 422
    # 422: over the store cap after cleaning (store-enforced)...
    too_long = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "x" * 4001}
    )
    assert too_long.status_code == 422
    # ...and far over the boundary bound (Pydantic-enforced, defense in depth).
    huge = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "x" * 6000}
    )
    assert huge.status_code == 422
    # None of the rejects persisted anything.
    assert auth_client.get(f"/whitepaper/runs/{run_id}").json()["inputs"] == {}


# ---------------------------------------------------------------------------
# finalize / reopen lifecycle + audit rows
# ---------------------------------------------------------------------------


def test_finalize_freezes_reopen_unfreezes_and_audits(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    saved = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "Site A"}
    )
    assert saved.status_code == 200

    final = auth_client.post(f"/whitepaper/runs/{run_id}/finalize")
    assert final.status_code == 200
    assert final.json() == {"run_id": run_id, "status": "final"}

    # Final freezes the analyst layer: edit, clear, re-finalize, delete all 409.
    frozen = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "late"}
    )
    assert frozen.status_code == 409
    frozen_clear = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": None}
    )
    assert frozen_clear.status_code == 409
    assert auth_client.post(f"/whitepaper/runs/{run_id}/finalize").status_code == 409
    assert auth_client.delete(f"/whitepaper/runs/{run_id}").status_code == 409
    # The frozen overlay survived the rejected edits.
    detail = auth_client.get(f"/whitepaper/runs/{run_id}").json()
    assert detail["status"] == "final"
    assert detail["finalized_by"] == "Test Analyst"
    assert detail["finalized_at"] is not None
    assert detail["inputs"][ANALYST_CELL]["value"] == "Site A"

    reopened = auth_client.post(f"/whitepaper/runs/{run_id}/reopen")
    assert reopened.status_code == 200
    assert reopened.json() == {"run_id": run_id, "status": "draft"}
    # Reopen-then-edit works again.
    again = auth_client.post(
        f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "Site B"}
    )
    assert again.status_code == 200

    # Both workflow actions wrote audit rows on the shared QueryLog trail.
    rows = _whitepaper_audit_rows()
    assert [r["status"] for r in rows] == ["populated", "finalized", "reopened"]
    for row in rows[1:]:
        assert row["model_name"] == "(workflow)"
        assert row["route_json"]["route"] == "whitepaper"
        assert row["route_json"]["run_id"] == run_id
        assert row["route_json"]["application_number"] == APPL_NO
    assert rows[1]["route_json"]["reason"] == "finalized"
    assert rows[2]["route_json"]["reason"] == "reopened"


def test_finalize_audit_write_failure_keeps_the_run_final_and_does_not_500(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """INV-6/INV-4: the audit write runs AFTER finalize_run committed, so it must
    have a DEFINED failure. A naked 500 here would tell the analyst the finalize
    failed while the DB holds ``final`` -- and the retry 409s, so no later request
    could ever write the row either."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    monkeypatch.setattr(api_main, "log_query", _boom)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/finalize")
    assert r.status_code == 200, r.text
    assert r.json() == {"run_id": run_id, "status": "final"}
    # The response told the truth about the committed state transition.
    assert auth_client.get(f"/whitepaper/runs/{run_id}").json()["status"] == "final"
    # Pin the LOUD degrade: the row is missing (logged + Sentry-captured), not
    # silently invented.
    assert [row["status"] for row in _whitepaper_audit_rows()] == ["populated"]


def test_reopen_audit_write_failure_keeps_the_run_draft_and_does_not_500(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The reopen twin of the finalize case: reopen_run has already committed
    draft, and the retry 409s on RunNotFinalError."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    assert auth_client.post(f"/whitepaper/runs/{run_id}/finalize").status_code == 200
    monkeypatch.setattr(api_main, "log_query", _boom)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/reopen")
    assert r.status_code == 200, r.text
    assert r.json() == {"run_id": run_id, "status": "draft"}
    assert auth_client.get(f"/whitepaper/runs/{run_id}").json()["status"] == "draft"
    assert [row["status"] for row in _whitepaper_audit_rows()] == ["populated", "finalized"]


def test_finalize_appl_no_lookup_failure_still_writes_the_audit_row(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A blip on the cosmetic application_number round-trip must not cost the
    audit row: the label degrades to empty, the row still lands.

    Only main's ``session_scope`` name is rebound, so the audit insert (which
    goes through ``regwatch.common.audit``'s own import) still succeeds -- that
    is the transient-fault shape this recovers.
    """
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    monkeypatch.setattr(api_main, "session_scope", _boom)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/finalize")
    assert r.status_code == 200, r.text
    rows = _whitepaper_audit_rows()
    assert [row["status"] for row in rows] == ["populated", "finalized"]
    assert rows[1]["route_json"]["run_id"] == run_id
    assert rows[1]["route_json"]["application_number"] == ""
    assert rows[1]["route_json"]["reason"] == "finalized"


def test_reopen_draft_is_409(auth_client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    assert auth_client.post(f"/whitepaper/runs/{run_id}/reopen").status_code == 409


# ---------------------------------------------------------------------------
# delete rules
# ---------------------------------------------------------------------------


def test_delete_creator_only_and_drafts_only(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    other = _other_client()
    try:
        # Foreign delete is 403 even on a draft; ownership outranks state.
        assert other.delete(f"/whitepaper/runs/{run_id}").status_code == 403
    finally:
        other.__exit__(None, None, None)
    r = auth_client.delete(f"/whitepaper/runs/{run_id}")
    assert r.status_code == 200
    assert r.json() == {"deleted": True, "run_id": run_id}
    assert auth_client.get(f"/whitepaper/runs/{run_id}").status_code == 404
    assert auth_client.get("/whitepaper/runs").json()["total"] == 0


def test_foreign_delete_of_final_run_is_403_not_409(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Ownership is checked before state: a non-creator never learns more than
    'not yours' about a finalized run's lifecycle."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    assert auth_client.post(f"/whitepaper/runs/{run_id}/finalize").status_code == 200
    other = _other_client()
    try:
        assert other.delete(f"/whitepaper/runs/{run_id}").status_code == 403
    finally:
        other.__exit__(None, None, None)


# ---------------------------------------------------------------------------
# server-side docx render (from storage, with the overlay)
# ---------------------------------------------------------------------------


def test_docx_renders_from_storage_with_overlay(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    install_fake_sources(monkeypatch)
    body = _populate(auth_client)
    run_id = body["run_id"]
    assert (
        auth_client.post(
            f"/whitepaper/runs/{run_id}/cells/{ANALYST_CELL}", json={"value": "Site A approved"}
        ).status_code
        == 200
    )
    assert (
        auth_client.post(
            f"/whitepaper/runs/{run_id}/cells/{POPULATED_CELL}", json={"value": "verify brand"}
        ).status_code
        == 200
    )
    _forbid_repopulate(monkeypatch)
    _use_absent_template(monkeypatch, tmp_path)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/docx")
    assert r.status_code == 200, r.text
    assert "wordprocessingml.document" in r.headers["content-type"]
    assert f"whitepaper_{APPL_NO}.docx" in r.headers["content-disposition"]
    assert r.content[:2] == b"PK"
    doc = Document(BytesIO(r.content))
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n".join(
        c.text for t in doc.tables for row in t.rows for c in row.cells
    )
    # The analyst layer rendered attributed; the generated value survived verbatim.
    assert "Site A approved [analyst: Test Analyst]" in all_text
    assert "Analyst note (Test Analyst): verify brand" in all_text
    assert "ALBUTEROL SULFATE" in all_text
    assert "Analyst inputs" in {p.text for p in doc.paragraphs}

    # One lightweight docx_rendered audit row rides on the populate row.
    rows = _whitepaper_audit_rows()
    assert [row["status"] for row in rows] == ["populated", "docx_rendered"]
    render_row = rows[1]
    assert render_row["model_name"] == "(docx-render)"
    assert render_row["route_json"]["reason"] == "docx_render"
    assert render_row["route_json"]["run_id"] == run_id
    assert render_row["route_json"]["source_audit_id"] == body["audit_id"]
    assert APPL_NO in render_row["query_text"]


def test_docx_stored_corruption_is_500_and_no_document(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    """Tampered stored sections no longer hash to sections_sha256: 500, no
    document, no docx audit row -- and finalize refuses the same way."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    _forbid_repopulate(monkeypatch)
    _use_absent_template(monkeypatch, tmp_path)
    with session_scope() as s:
        run = s.get(WhitepaperRun, run_id)
        assert run is not None
        tampered = copy.deepcopy(run.sections_json)
        tampered[0]["cells"][0]["value"] = "TAMPERED"
        run.sections_json = tampered
        s.add(run)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/docx")
    assert r.status_code == 500
    assert "integrity" in r.json()["detail"]
    assert "content-disposition" not in {k.lower() for k in r.headers}
    assert auth_client.post(f"/whitepaper/runs/{run_id}/finalize").status_code == 500
    # No render/finalize audit rows were written; the run stays draft.
    assert [row["status"] for row in _whitepaper_audit_rows()] == ["populated"]
    assert auth_client.get(f"/whitepaper/runs/{run_id}").json()["status"] == "draft"


def test_docx_header_unsafe_stored_application_number_is_500(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    """The stored application number is interpolated into Content-Disposition:
    anything not application-number-shaped is refused (ported filename guard)."""
    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    _forbid_repopulate(monkeypatch)
    _use_absent_template(monkeypatch, tmp_path)
    with session_scope() as s:
        run = s.get(WhitepaperRun, run_id)
        assert run is not None
        run.application_number = '0205"03\r\nX-Evil: injected'
        s.add(run)

    r = auth_client.post(f"/whitepaper/runs/{run_id}/docx")
    assert r.status_code == 500
    assert "x-evil" not in {k.lower() for k in r.headers}
    assert "content-disposition" not in {k.lower() for k in r.headers}
    assert [row["status"] for row in _whitepaper_audit_rows()] == ["populated"]


def test_docx_rate_limited(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    """The render keeps drawing from the /query budget (CPU-bound assembly)."""
    import config.settings as cs

    install_fake_sources(monkeypatch)
    run_id = _populate(auth_client)["run_id"]
    _forbid_repopulate(monkeypatch)
    _use_absent_template(monkeypatch, tmp_path)

    monkeypatch.setenv("RATE_LIMIT_PER_MINUTE", "1")
    cs.get_settings.cache_clear()
    try:
        assert auth_client.post(f"/whitepaper/runs/{run_id}/docx").status_code == 200
        assert auth_client.post(f"/whitepaper/runs/{run_id}/docx").status_code == 429
    finally:
        cs.get_settings.cache_clear()


# ---------------------------------------------------------------------------
# /health template observability (design doc section 8)
# ---------------------------------------------------------------------------


def test_health_reports_whitepaper_template_status(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    """A prod stack silently rendering FALLBACK_MARKER documents must be visible:
    /health (open, no auth) names the template state without any fetch."""
    import config.settings as cs

    client = TestClient(app)
    client.__enter__()
    try:
        monkeypatch.setenv("WHITEPAPER_TEMPLATE_PATH", str(tmp_path / "absent.docx"))
        monkeypatch.delenv("WHITEPAPER_TEMPLATE_URL", raising=False)
        cs.get_settings.cache_clear()
        assert client.get("/health").json()["whitepaper_template"] == "absent"

        monkeypatch.setenv("WHITEPAPER_TEMPLATE_URL", "https://example.invalid/signed-url")
        cs.get_settings.cache_clear()
        assert client.get("/health").json()["whitepaper_template"] == "fetchable"

        present = tmp_path / "present.docx"
        present.write_bytes(b"PK\x03\x04")
        monkeypatch.setenv("WHITEPAPER_TEMPLATE_PATH", str(present))
        cs.get_settings.cache_clear()
        assert client.get("/health").json()["whitepaper_template"] == "present"
    finally:
        cs.get_settings.cache_clear()
        client.__exit__(None, None, None)
