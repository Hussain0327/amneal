"""Reference-library API tests: /psg/documents listing, PDF, text and docx.

The listing feeds the Compliance Studio rail; the PDF route streams the
document inline; the content and docx routes render the same document as
studio blocks and as a Word download. These cover the failure paths the routes
exist to make legible: auth, pagination determinism, local-cache hits vs the
remote fda.gov branch, the ETag short-circuit, every mapped upstream failure,
and the two ways a document can exist without renderable text.
"""

from __future__ import annotations

import hashlib
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

import httpx
import pytest
import respx
from config.settings import get_settings
from docx import Document
from fastapi.testclient import TestClient

from regwatch.api.main import app
from regwatch.store.db import session_scope
from regwatch.store.models import PsgDocument, PsgVersion
from regwatch.store.vector_store import add_chunks

_PDF = b"%PDF-1.4 tiny reference-library test body"
_PDF_HASH = hashlib.sha256(_PDF).hexdigest()
_FDA_URL = "https://www.accessdata.fda.gov/drugsatfda_docs/psg/PSG_020503.pdf"


def _mk_doc(**overrides: Any) -> int:
    """Insert one psg_document row directly; returns its id."""
    fields: dict[str, Any] = {
        "active_ingredient": "Albuterol Sulfate",
        "normalized_name": "albuterol sulfate",
        "dosage_form": "Aerosol, Metered",
        "route": "Inhalation",
        "appl_no": None,
        "psg_type": "final",
        "recommended_date": "2024-05-01",
        "source_url": _FDA_URL,
        "pdf_path": None,
        "content_hash": _PDF_HASH,
    }
    fields.update(overrides)
    with session_scope() as s:
        row = PsgDocument(**fields)
        s.add(row)
        s.flush()
        assert row.id is not None
        return row.id


# ---------------------------------------------------------------------------
# GET /psg/documents (listing)
# ---------------------------------------------------------------------------


def test_list_requires_auth_401() -> None:
    with TestClient(app) as client:
        assert client.get("/psg/documents").status_code == 401


def test_list_empty_db(auth_client: TestClient) -> None:
    body = auth_client.get("/psg/documents").json()
    assert body == {"count": 0, "total": 0, "limit": 2000, "offset": 0, "documents": []}


def test_list_fields_and_stripped_name(auth_client: TestClient) -> None:
    doc_id = _mk_doc(appl_no="020503")
    body = auth_client.get("/psg/documents").json()
    assert body["count"] == body["total"] == 1
    assert body["documents"] == [
        {
            "id": doc_id,
            "active_ingredient": "Albuterol Sulfate",
            "normalized_name": "albuterol sulfate",
            "stripped_name": "albuterol",
            "dosage_form": "Aerosol, Metered",
            "route": "Inhalation",
            "appl_no": "020503",
            "psg_type": "final",
            "recommended_date": "2024-05-01",
            "source_url": _FDA_URL,
        }
    ]


def test_list_ordering_deterministic(auth_client: TestClient) -> None:
    # Inserted deliberately out of order; the response must sort by
    # normalized_name, then NULL-as-"" dosage_form, then route, then psg_type.
    d_b = _mk_doc(normalized_name="budesonide", active_ingredient="Budesonide")
    d_a_tab_final = _mk_doc(normalized_name="albuterol sulfate", dosage_form="Tablet")
    d_a_null = _mk_doc(normalized_name="albuterol sulfate", dosage_form=None)
    d_a_tab_draft = _mk_doc(
        normalized_name="albuterol sulfate", dosage_form="Tablet", psg_type="draft"
    )
    body = auth_client.get("/psg/documents").json()
    got = [doc["id"] for doc in body["documents"]]
    # NULL dosage_form coalesces to "" (sorts first); draft < final on ties.
    assert got == [d_a_null, d_a_tab_draft, d_a_tab_final, d_b]


def test_list_pagination_gapless(auth_client: TestClient) -> None:
    ids = {
        _mk_doc(normalized_name=name, active_ingredient=name.title())
        for name in ("alpha", "bravo", "charlie")
    }
    page1 = auth_client.get("/psg/documents", params={"limit": 2, "offset": 0}).json()
    page2 = auth_client.get("/psg/documents", params={"limit": 2, "offset": 2}).json()
    assert page1["count"] == 2 and page1["total"] == 3
    assert page2["count"] == 1 and page2["total"] == 3
    union = [d["id"] for d in page1["documents"]] + [d["id"] for d in page2["documents"]]
    assert len(union) == len(set(union)) == 3
    assert set(union) == ids


def test_list_limit_bounds_422(auth_client: TestClient) -> None:
    assert auth_client.get("/psg/documents", params={"limit": 5001}).status_code == 422
    assert auth_client.get("/psg/documents", params={"limit": 0}).status_code == 422


# ---------------------------------------------------------------------------
# GET/HEAD /psg/documents/{doc_id}/pdf
# ---------------------------------------------------------------------------


def test_pdf_requires_auth_401() -> None:
    with TestClient(app) as client:
        assert client.get("/psg/documents/1/pdf").status_code == 401
        assert client.head("/psg/documents/1/pdf").status_code == 401


def test_pdf_unknown_id_404(auth_client: TestClient) -> None:
    assert auth_client.get("/psg/documents/999999/pdf").status_code == 404


@respx.mock  # no routes mocked: ANY network call raises and fails the test
def test_pdf_local_hit_via_pdf_path(auth_client: TestClient, tmp_path: Any) -> None:
    local = tmp_path / "stored.pdf"
    local.write_bytes(_PDF)
    doc_id = _mk_doc(appl_no="020503", pdf_path=str(local))
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert resp.headers["content-disposition"] == 'inline; filename="PSG_020503.pdf"'
    assert resp.headers["etag"] == f'"{_PDF_HASH}"'
    assert resp.content == _PDF


@respx.mock
def test_pdf_local_hit_via_deterministic_cache_path(auth_client: TestClient) -> None:
    raw_dir = get_settings().raw_pdf_dir
    raw_dir.mkdir(parents=True, exist_ok=True)
    (raw_dir / f"PSG_020503_{_PDF_HASH[:8]}.pdf").write_bytes(_PDF)
    doc_id = _mk_doc(appl_no="020503", pdf_path=None)
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    assert resp.content == _PDF


@respx.mock
def test_pdf_etag_304_short_circuits(auth_client: TestClient, tmp_path: Any) -> None:
    # pdf_path points nowhere and raw_pdf_dir holds no file: only the DB row
    # can answer, so a 304 here proves neither disk nor network was consulted
    # (respx with zero routes raises on any request).
    doc_id = _mk_doc(appl_no="020503", pdf_path=str(tmp_path / "gone.pdf"))
    for header_value in (f'"{_PDF_HASH}"', f'W/"{_PDF_HASH}"', f'"other", "{_PDF_HASH}"'):
        resp = auth_client.get(
            f"/psg/documents/{doc_id}/pdf", headers={"If-None-Match": header_value}
        )
        assert resp.status_code == 304
        assert resp.headers["etag"] == f'"{_PDF_HASH}"'


@respx.mock
def test_pdf_remote_fetch_serves_and_caches(auth_client: TestClient) -> None:
    respx.get(_FDA_URL).mock(return_value=httpx.Response(200, content=_PDF))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    assert resp.content == _PDF
    assert resp.headers["etag"] == f'"{_PDF_HASH}"'
    # download_pdf's write-through cache: the next open on this machine is a
    # local hit under the crawler's own naming scheme.
    assert (get_settings().raw_pdf_dir / f"PSG_020503_{_PDF_HASH[:8]}.pdf").is_file()


@respx.mock
def test_pdf_remote_timeout_504(auth_client: TestClient) -> None:
    respx.get(_FDA_URL).mock(side_effect=httpx.ConnectTimeout("boom"))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 504
    assert "fda.gov" in resp.json()["detail"]
    assert _FDA_URL not in resp.json()["detail"]


@respx.mock
def test_pdf_remote_non_pdf_502(auth_client: TestClient) -> None:
    respx.get(_FDA_URL).mock(
        return_value=httpx.Response(200, content=b"<html>challenge page</html>")
    )
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "upstream did not return a PDF"


@respx.mock
def test_pdf_remote_oversize_502(auth_client: TestClient) -> None:
    # Oversize by declared Content-Length: exercises the cap without a 60MB body.
    respx.get(_FDA_URL).mock(
        return_value=httpx.Response(
            200, headers={"content-length": str(60_000_000)}, content=b"%PDF-"
        )
    )
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "upstream PDF exceeds the size cap"


@respx.mock
def test_pdf_upstream_404_502_without_url(auth_client: TestClient) -> None:
    route = respx.get(_FDA_URL).mock(return_value=httpx.Response(404))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert "404" in resp.json()["detail"]
    assert _FDA_URL not in resp.json()["detail"]
    assert route.call_count == 1  # 4xx is terminal, never retried


@respx.mock
def test_pdf_retried_out_5xx_502(auth_client: TestClient) -> None:
    # A persistent 503 exhausts the crawler's bounded retry and reraises its
    # INTERNAL marker class, not HTTPStatusError -- the route must still map
    # it to a legible 502 instead of a 500.
    route = respx.get(_FDA_URL).mock(return_value=httpx.Response(503))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "fda.gov kept failing to serve the PDF"
    assert route.call_count == 3  # bounded retry, not a storm


@respx.mock
def test_pdf_remote_redirect_loop_502(auth_client: TestClient) -> None:
    # TooManyRedirects subclasses RequestError, NOT TransportError -- the
    # catch-all arm must be httpx.HTTPError or this escapes as a 500.
    respx.get(_FDA_URL).mock(side_effect=httpx.TooManyRedirects("loop"))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "could not reach fda.gov for the PDF"


@respx.mock
def test_pdf_remote_decoding_error_502(auth_client: TestClient) -> None:
    respx.get(_FDA_URL).mock(side_effect=httpx.DecodingError("bad gzip"))
    doc_id = _mk_doc(appl_no="020503")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "could not reach fda.gov for the PDF"


@respx.mock
def test_pdf_unwritable_cache_dir_still_serves(auth_client: TestClient) -> None:
    # A full/read-only disk must not turn a successful fetch into a 500: the
    # write-through cache is best-effort, the bytes in hand are the answer.
    raw_dir = get_settings().raw_pdf_dir
    raw_dir.mkdir(parents=True, exist_ok=True)
    raw_dir.chmod(0o500)
    try:
        respx.get(_FDA_URL).mock(return_value=httpx.Response(200, content=_PDF))
        doc_id = _mk_doc(appl_no="020503")
        resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
        assert resp.status_code == 200
        assert resp.content == _PDF
    finally:
        raw_dir.chmod(0o700)


@respx.mock
def test_pdf_corrupt_local_file_falls_through_to_remote(
    auth_client: TestClient, tmp_path: Any
) -> None:
    # A truncated/foreign file at the expected path must never be served under
    # the content_hash ETag; the verified miss recovers via fda.gov.
    bad = tmp_path / "truncated.pdf"
    bad.write_bytes(_PDF[: len(_PDF) // 2])
    respx.get(_FDA_URL).mock(return_value=httpx.Response(200, content=_PDF))
    doc_id = _mk_doc(appl_no="020503", pdf_path=str(bad))
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    assert resp.content == _PDF
    assert resp.headers["etag"] == f'"{_PDF_HASH}"'


@respx.mock
def test_pdf_hash_drift_serves_with_new_etag(auth_client: TestClient) -> None:
    respx.get(_FDA_URL).mock(return_value=httpx.Response(200, content=_PDF))
    doc_id = _mk_doc(appl_no="020503", content_hash="0" * 64)
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    # FDA revised since the row was written: the served ETag must describe the
    # bytes actually returned, not the stale row.
    assert resp.headers["etag"] == f'"{_PDF_HASH}"'


@respx.mock
def test_pdf_non_fda_source_url_502(auth_client: TestClient) -> None:
    doc_id = _mk_doc(appl_no="020503", source_url="https://evil.example/x.pdf")
    resp = auth_client.get(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 502
    assert resp.json()["detail"] == "document has no fetchable FDA source"


@respx.mock
def test_pdf_remote_rate_limited_429_local_hits_unmetered(
    auth_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Any
) -> None:
    import config.settings as cs

    monkeypatch.setenv("RATE_LIMIT_PER_MINUTE", "1")
    cs.get_settings.cache_clear()
    respx.get(_FDA_URL).mock(return_value=httpx.Response(200, content=_PDF))
    remote_id = _mk_doc(appl_no="020503")
    assert auth_client.get(f"/psg/documents/{remote_id}/pdf").status_code == 200
    # Force the second request onto the remote branch again by removing the
    # write-through cache file the first request created.
    cached = get_settings().raw_pdf_dir / f"PSG_020503_{_PDF_HASH[:8]}.pdf"
    cached.unlink()
    assert auth_client.get(f"/psg/documents/{remote_id}/pdf").status_code == 429
    # A local-disk hit does not touch the remote budget.
    local = tmp_path / "local.pdf"
    local.write_bytes(_PDF)
    local_id = _mk_doc(appl_no="020504", pdf_path=str(local))
    assert auth_client.get(f"/psg/documents/{local_id}/pdf").status_code == 200


@respx.mock
def test_pdf_head_answers_from_db_row_only(auth_client: TestClient, tmp_path: Any) -> None:
    # No local file, no mocked network route: a 200 proves the probe never
    # leaves the DB row (an FDA source_url is enough to vouch for the GET).
    doc_id = _mk_doc(appl_no="020503", pdf_path=str(tmp_path / "gone.pdf"))
    resp = auth_client.head(f"/psg/documents/{doc_id}/pdf")
    assert resp.status_code == 200
    assert resp.headers["etag"] == f'"{_PDF_HASH}"'
    assert resp.headers["content-disposition"] == 'inline; filename="PSG_020503.pdf"'
    assert resp.content == b""
    assert auth_client.head("/psg/documents/999999/pdf").status_code == 404


@respx.mock
def test_pdf_head_refuses_guaranteed_fail_row(auth_client: TestClient, tmp_path: Any) -> None:
    # Non-FDA source_url and no local file: the GET would 502 with certainty,
    # so a 200 probe would defeat the probe. A local file rescues it.
    doc_id = _mk_doc(appl_no="020503", source_url="https://evil.example/x.pdf")
    assert auth_client.head(f"/psg/documents/{doc_id}/pdf").status_code == 502
    local = tmp_path / "cached.pdf"
    local.write_bytes(_PDF)
    rescued = _mk_doc(
        appl_no="020504", source_url="https://evil.example/y.pdf", pdf_path=str(local)
    )
    assert auth_client.head(f"/psg/documents/{rescued}/pdf").status_code == 200


# ---------------------------------------------------------------------------
# GET /psg/documents/{id}/content and /docx
# ---------------------------------------------------------------------------

_PAGE_1 = (
    "Draft Guidance on Albuterol Sulfate\n"
    "May 2026\n"
    "Active Ingredient: Albuterol sulfate\n"
    "Recommended Studies:\n"
    "Three in vitro bioequivalence studies are recommended for this product."
)
_PAGE_2 = "Document History: Recommended May 2026"


def _mk_version(doc_id: int, *, content_hash: str = _PDF_HASH) -> int:
    """Insert one psg_version row for a document; returns its id."""
    with session_scope() as s:
        row = PsgVersion(
            psg_document_id=doc_id,
            content_hash=content_hash,
            recommended_date="2024-05-01",
            captured_at=datetime.now(UTC).replace(tzinfo=None),
        )
        s.add(row)
        s.flush()
        assert row.id is not None
        return row.id


def _mk_text(doc_id: int, version_id: int, *pages: str) -> None:
    """Store one chunk per page for a version, the way ingest would.

    Embeddings are None on purpose: this route reads chunks, it never searches
    them, so the test does not need an embedding provider to exercise it.
    """
    texts = list(pages)
    add_chunks(
        ids=[f"{doc_id}-{version_id}-{i}" for i in range(len(texts))],
        embeddings=[None] * len(texts),
        documents=texts,
        metadatas=[
            {
                "doc_id": doc_id,
                "version_id": version_id,
                "ordinal": i,
                "page": i + 1,
                "normalized_name": "albuterol sulfate",
                "appl_no": "020503",
                "source_url": _FDA_URL,
                "section_path": "",
            }
            for i in range(len(texts))
        ],
    )


def _seeded_doc() -> int:
    """A document with a current version and its text. Returns the doc id."""
    doc_id = _mk_doc(appl_no="020503")
    version_id = _mk_version(doc_id)
    _mk_text(doc_id, version_id, _PAGE_1, _PAGE_2)
    return doc_id


def test_content_requires_auth_401() -> None:
    with TestClient(app) as client:
        assert client.get("/psg/documents/1/content").status_code == 401


def test_docx_requires_auth_401() -> None:
    with TestClient(app) as client:
        assert client.get("/psg/documents/1/docx").status_code == 401


def test_content_unknown_id_404(auth_client: TestClient) -> None:
    assert auth_client.get("/psg/documents/9999/content").status_code == 404


def test_content_document_without_a_version_404(auth_client: TestClient) -> None:
    doc_id = _mk_doc(appl_no="020503")
    res = auth_client.get(f"/psg/documents/{doc_id}/content")
    assert res.status_code == 404
    assert "no stored version" in res.json()["detail"]


def test_content_version_without_text_409(auth_client: TestClient) -> None:
    # The row is real and the PDF is still servable, so this is not a 404: the
    # client should offer the PDF rather than report a bad id.
    doc_id = _mk_doc(appl_no="020503")
    _mk_version(doc_id)
    res = auth_client.get(f"/psg/documents/{doc_id}/content")
    assert res.status_code == 409
    assert "not available" in res.json()["detail"]


def test_content_returns_the_document_as_blocks(auth_client: TestClient) -> None:
    doc_id = _seeded_doc()
    body = auth_client.get(f"/psg/documents/{doc_id}/content").json()

    assert body["id"] == doc_id
    assert body["appl_no"] == "020503"
    assert body["file_name"] == "PSG_020503 Albuterol Sulfate.docx"
    assert body["active_ingredient"] == "Albuterol Sulfate"
    assert body["dosage_form"] == "Aerosol, Metered"
    assert body["route"] == "Inhalation"
    assert body["psg_type"] == "final"
    assert body["recommended_date"] == "2024-05-01"
    assert body["source_url"] == _FDA_URL
    assert body["page_count"] == 2
    assert body["truncated"] is False

    kinds = [(b["type"], b["text"]) for b in body["blocks"]]
    assert kinds[0] == ("title", "Draft Guidance on Albuterol Sulfate")
    assert ("h2", "Recommended Studies:") in kinds
    assert ("meta", "Document History: Recommended May 2026") in kinds
    assert all(b["id"].startswith(f"psg-{doc_id}-b") for b in body["blocks"])
    assert [b["page"] for b in body["blocks"]][-1] == 2


def test_content_renders_only_the_current_version(auth_client: TestClient) -> None:
    # A revision's text supersedes the previous version's in the index, but a
    # row for the old version can still exist; the studio must never show it.
    doc_id = _mk_doc(appl_no="020503")
    old = _mk_version(doc_id, content_hash="a" * 64)
    _mk_text(doc_id, old, "Title line\nSuperseded guidance text.")
    current = _mk_version(doc_id, content_hash="b" * 64)
    _mk_text(doc_id, current, "Title line\nCurrent guidance text.")

    body = auth_client.get(f"/psg/documents/{doc_id}/content").json()
    rendered = " ".join(b["text"] for b in body["blocks"])
    assert "Current guidance text." in rendered
    assert "Superseded" not in rendered


def test_docx_downloads_with_a_safe_filename(auth_client: TestClient) -> None:
    doc_id = _seeded_doc()
    res = auth_client.get(f"/psg/documents/{doc_id}/docx")

    assert res.status_code == 200
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml"
    )
    assert res.headers["content-disposition"] == (
        'attachment; filename="PSG_020503_Albuterol_Sulfate.docx"'
    )
    # A real OOXML package, not an error body with a hopeful content type.
    assert res.content[:2] == b"PK"


def test_docx_carries_the_same_text_the_content_route_returns(
    auth_client: TestClient,
) -> None:
    doc_id = _seeded_doc()
    blocks = auth_client.get(f"/psg/documents/{doc_id}/content").json()["blocks"]
    data = auth_client.get(f"/psg/documents/{doc_id}/docx").content

    document = Document(BytesIO(data))
    rendered = [p.text for p in document.paragraphs]
    for block in blocks:
        assert block["text"] in rendered


def test_docx_unknown_id_404(auth_client: TestClient) -> None:
    assert auth_client.get("/psg/documents/9999/docx").status_code == 404


def test_docx_filename_survives_a_punctuated_combination_product(
    auth_client: TestClient,
) -> None:
    doc_id = _mk_doc(
        appl_no=None,
        active_ingredient="Ethinyl Estradiol; Levonorgestrel",
        normalized_name="ethinyl estradiol; levonorgestrel",
    )
    version_id = _mk_version(doc_id)
    _mk_text(doc_id, version_id, _PAGE_1)

    res = auth_client.get(f"/psg/documents/{doc_id}/docx")
    assert res.headers["content-disposition"] == (
        'attachment; filename="Ethinyl_Estradiol_Levonorgestrel_PSG.docx"'
    )
