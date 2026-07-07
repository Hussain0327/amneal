"""White-Paper docx writer — synthetic template fill + from-scratch fallback."""

from __future__ import annotations

import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document

import regwatch.whitepaper.docx_writer as docx_writer
from regwatch.whitepaper.docx_writer import (
    FALLBACK_MARKER,
    _marker,
    _render_value,
    docx_media_type,
    write_whitepaper_docx,
)
from regwatch.whitepaper.populator import build_whitepaper
from tests._whitepaper_stub import APPL_NO, RLD_NAME, install_fake_sources


def _synthetic_template(path: Path) -> None:
    """A tiny structurally-equivalent template (label | value 2-col tables)."""
    doc = Document()
    doc.add_heading("CRA White Paper", level=0)
    sec1 = doc.add_table(rows=0, cols=2)
    for label in ("Product Name", "Strengths", "Reference Listed Drug (RLD)"):
        row = sec1.add_row().cells
        row[0].text = label
        row[1].text = ""
    checks = doc.add_table(rows=0, cols=2)
    for label in ("REMS", "Drug Shortages", "DEA Classification"):
        row = checks.add_row().cells
        row[0].text = label
        row[1].text = "Yes\t\tNo"
    doc.save(str(path))


def _build(monkeypatch: pytest.MonkeyPatch) -> dict:
    install_fake_sources(monkeypatch)
    return build_whitepaper(RLD_NAME, APPL_NO)


def test_from_scratch_when_template_absent(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=tmp_path / "does-not-exist.docx")
    assert data[:2] == b"PK"  # a real .docx (zip) was produced
    doc = Document(BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "CRA White Paper" in text
    # All four section headings + the provenance appendix render.
    headings = {p.text for p in doc.paragraphs}
    assert "Provenance appendix" in headings
    # The populated product name and a structured locator are present somewhere.
    all_cells = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "ALBUTEROL SULFATE" in all_cells
    assert "OB_020503/001" in all_cells


def test_fills_real_template(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    _synthetic_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=template)
    doc = Document(BytesIO(data))
    cell_text = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    joined = "\n".join(cell_text)
    # Value cells filled in place.
    assert any("ALBUTEROL SULFATE" in c for c in cell_text)
    # Checkbox rows get an appended marker rather than overwriting the Yes/No text.
    assert "->" in joined
    assert "Analyst input required" in joined  # DEA collapses -> marker
    # Provenance appendix present.
    assert "Provenance appendix" in "\n".join(p.text for p in doc.paragraphs)


def _real_shaped_template(path: Path) -> None:
    """Mirrors the real template's hard rows: the merged Priority-Status
    checkbox block, parenthetical label suffixes, the merged "Dosage Form;
    Route" row, and a duplicated label row."""
    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    for label in (
        "Product Name",
        "Dosage Form; Route",
        "Established Pharmacologic Class (A scientifically valid classification system "
        "used by the FDA in drug labeling)",
        "Pregnancy Registry Contact Detail (if required)",
        "Salable Unit (from Sales & Marketing)",
    ):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = ""
    # The real template repeats the Combination Product label (Yes/No row +
    # the Type 1-9 list row).
    for prefill in ("No / Yes", "Type 1 / Type 2 / Type 3"):
        row = table.add_row().cells
        row[0].text = "Combination Product"
        row[1].text = prefill
    # Merged Priority Status block: one value cell carrying the whole
    # Patents / PI-PIV / eFTF / Drug-Shortages checkbox section.
    row = table.add_row().cells
    row[0].text = "Priority Status"
    block = row[1]
    block.text = "Patents (OB review):"
    block.add_paragraph("No Relevant Patents / PI / PII / PIII / PIV / Section viii")
    block.add_paragraph("If PI - Eligible for First to Market? Yes/No")
    block.add_paragraph("If PIV - Eligible for eFTF (PIV Website review)? Yes/No")
    block.add_paragraph("Drug Shortages:")
    block.add_paragraph("On Drug Shortage List? Yes/No")
    doc.save(str(path))


def _table_cell_texts(data: bytes) -> list[str]:
    doc = Document(BytesIO(data))
    return [c.text for t in doc.tables for r in t.rows for c in r.cells]


def test_priority_block_preserved_and_marked(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    template = tmp_path / "real-shaped.docx"
    _real_shaped_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=template)
    doc = Document(BytesIO(data))
    block_cell = next(
        r.cells[-1]
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip() == "Priority Status"
    )
    # The template's own checkbox lines survive...
    assert "If PIV - Eligible for eFTF (PIV Website review)? Yes/No" in block_cell.text
    assert "On Drug Shortage List? Yes/No" in block_cell.text
    # ...and the block members ride in as appended markers (drug_shortage is
    # verified_absent in the stub -> "No"; the manual cells -> analyst marker).
    assert "Drug Shortages: On Shortage List? Y/N  ->  No" in block_cell.text
    assert "Analyst input required" in block_cell.text


def test_prefix_matched_labels_fill_their_rows(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    template = tmp_path / "real-shaped.docx"
    _real_shaped_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=template)
    doc = Document(BytesIO(data))
    by_label = {
        r.cells[0].text.strip(): r.cells[-1].text for t in doc.tables for r in t.rows if r.cells
    }
    epc_row = next(v for k, v in by_label.items() if k.startswith("Established Pharmacologic"))
    assert "[EPC]" in epc_row
    pregnancy_row = next(v for k, v in by_label.items() if k.startswith("Pregnancy Registry"))
    assert "1-800-555-0100" in pregnancy_row
    salable_row = next(v for k, v in by_label.items() if k.startswith("Salable Unit"))
    assert "Analyst input required" in salable_row
    # Route rides into the merged "Dosage Form; Route" row.
    dosage_row = by_label["Dosage Form; Route"]
    assert "AEROSOL, METERED" in dosage_row
    assert "Route: INHALATION" in dosage_row


def test_duplicate_label_rows_filled_once(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    template = tmp_path / "real-shaped.docx"
    _real_shaped_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=template)
    doc = Document(BytesIO(data))
    combo_rows = [
        r.cells[-1].text
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip() == "Combination Product"
    ]
    assert len(combo_rows) == 2
    marked = [text for text in combo_rows if "->" in text]
    assert len(marked) == 1  # the marker lands once, in the first matching row
    assert "Type 1 / Type 2 / Type 3" in combo_rows[1]
    assert "->" not in combo_rows[1]


def test_media_type() -> None:
    assert "wordprocessingml.document" in docx_media_type()


# ---------- empty/whitespace value defense (contract C3, writer side) ----------
def _populated_cell(value: object) -> dict:
    return {
        "id": "rems",
        "label": "REMS",
        "mode": "auto",
        "status": "populated",
        "value": value,
        "evidence": [],
        "note": None,
    }


def test_render_value_treats_empty_and_whitespace_as_no_value() -> None:
    assert _render_value(_populated_cell("Yes — program X")) == "Yes — program X"
    for blank in ("", "   ", "\n\t", None):
        assert _render_value(_populated_cell(blank)) == "Analyst input required"
    absent = {**_populated_cell(""), "status": "verified_absent"}
    assert _render_value(absent) == "No (verified absent)"


def test_marker_treats_empty_and_whitespace_as_no_value() -> None:
    assert _marker(_populated_cell("Yes — program X")) == "Yes — program X"
    # A populated checkbox cell with no usable value falls to the bare marker,
    # never the literal blank.
    for blank in ("", "   ", "\n\t", None):
        assert _marker(_populated_cell(blank)) == "Yes"
    assert _marker({**_populated_cell(""), "status": "verified_absent"}) == "No"
    assert (
        _marker({**_populated_cell(""), "status": "analyst_input_required"})
        == "Analyst input required"
    )


def _set_cell_value(result: dict, cell_id: str, value: str) -> str:
    for section in result["sections"]:
        for cell in section["cells"]:
            if cell["id"] == cell_id:
                cell["value"] = value
                return str(cell["label"])
    raise AssertionError(f"cell {cell_id} not in result")


def test_empty_values_render_as_analyst_input_in_document(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A blank that slips past the populator never renders as a populated cell."""
    result = _build(monkeypatch)
    empty_label = _set_cell_value(result, "product_name", "")
    whitespace_label = _set_cell_value(result, "strengths", "   ")
    data = write_whitepaper_docx(result, template_path=tmp_path / "missing.docx")
    doc = Document(BytesIO(data))
    # Only the section tables (header Cell | Value | Status) — the provenance
    # appendix repeats labels with source columns.
    by_label = {
        r.cells[0].text: r.cells[1].text
        for t in doc.tables
        if t.rows and [c.text for c in t.rows[0].cells][:3] == ["Cell", "Value", "Status"]
        for r in t.rows[1:]
    }
    assert by_label[empty_label] == "Analyst input required"
    assert by_label[whitespace_label] == "Analyst input required"


# ---------- template fallback is loud (logger + visible marker) ----------
class _LogRecorder:
    def __init__(self) -> None:
        self.events: list[tuple[str, dict[str, Any]]] = []

    def warning(self, event: str, **kwargs: Any) -> None:
        self.events.append((event, kwargs))


def test_missing_template_falls_back_with_warning_and_marker(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    result = _build(monkeypatch)
    recorder = _LogRecorder()
    monkeypatch.setattr(docx_writer, "log", recorder)
    missing = tmp_path / "official-template.docx"
    data = write_whitepaper_docx(result, template_path=missing)
    # The logger warning names the missed path...
    assert recorder.events, "no warning logged for the missing template"
    event, kwargs = recorder.events[0]
    assert event == "whitepaper_template_missing"
    assert kwargs["template_path"] == str(missing)
    # ...and the document itself carries the visible marker under the heading.
    doc = Document(BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    assert FALLBACK_MARKER in paragraphs
    assert paragraphs.index(FALLBACK_MARKER) == paragraphs.index("CRA White Paper") + 1


def test_real_template_render_has_no_fallback_marker(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    template = tmp_path / "template.docx"
    _synthetic_template(template)
    result = _build(monkeypatch)
    recorder = _LogRecorder()
    monkeypatch.setattr(docx_writer, "log", recorder)
    data = write_whitepaper_docx(result, template_path=template)
    assert recorder.events == []
    doc = Document(BytesIO(data))
    assert FALLBACK_MARKER not in {p.text for p in doc.paragraphs}


def test_provenance_appendix_lists_evidence(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    result = _build(monkeypatch)
    data = write_whitepaper_docx(result, template_path=tmp_path / "missing.docx")
    doc = Document(BytesIO(data))
    # Find the provenance table (4 columns: Cell, Source, Locator, Fetched at).
    prov = [t for t in doc.tables if t.rows and t.rows[0].cells[0].text == "Cell"]
    assert prov, "provenance table not found"
    rows_text = "\n".join(c.text for t in prov for r in t.rows for c in r.cells)
    assert "OB_020503/001" in rows_text
    assert "Orange Book" in rows_text


# ---------- analyst overlay (Phase 2: attributed inputs, INV-3 discipline) ----------
UPDATED_AT = "2026-07-07T12:00:00"


def _inputs(**values: str) -> dict[str, dict[str, str]]:
    return {
        cell_id: {"value": value, "author": "Second Analyst", "updated_at": UPDATED_AT}
        for cell_id, value in values.items()
    }


def _all_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    cells = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    return paragraphs + "\n" + cells


def _document_xml(data: bytes) -> bytes:
    """The document part only - zip metadata carries wall-clock timestamps."""
    with zipfile.ZipFile(BytesIO(data)) as archive:
        return archive.read("word/document.xml")


def test_overlay_fills_analyst_cell_with_attribution(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """An analyst_input_required cell renders the overlay text, visibly attributed."""
    template = tmp_path / "real-shaped.docx"
    _real_shaped_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result, template_path=template, inputs=_inputs(salable_unit="30-count blister pack")
    )
    doc = Document(BytesIO(data))
    # First match in document order = the template row (the analyst-inputs
    # appendix repeats the label later with different columns).
    salable_row = next(
        r.cells[-1].text
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip().startswith("Salable Unit")
    )
    assert "30-count blister pack [analyst: Second Analyst]" in salable_row
    assert "Analyst input required" not in salable_row


def test_overlay_on_populated_cell_appends_note_never_replaces(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A populated cell keeps its generated value verbatim; the human text rides
    in as a distinct note paragraph (INV-3)."""
    template = tmp_path / "template.docx"
    _synthetic_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result, template_path=template, inputs=_inputs(product_name="verify with brand team")
    )
    doc = Document(BytesIO(data))
    name_cell = next(
        r.cells[-1].text
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip() == "Product Name"
    )
    assert "ALBUTEROL SULFATE" in name_cell  # generated value untouched
    assert "Analyst note (Second Analyst): verify with brand team" in name_cell
    # The note never replaces the value: both live in the same cell, value first.
    assert name_cell.index("ALBUTEROL SULFATE") < name_cell.index("Analyst note")


def test_overlay_checkbox_cell_marker_format(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A checkbox-style analyst cell takes the overlay as its appended marker,
    keeping the template's own Yes/No text."""
    template = tmp_path / "template.docx"
    _synthetic_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result, template_path=template, inputs=_inputs(dea_classification="Not scheduled")
    )
    doc = Document(BytesIO(data))
    dea_cell = next(
        r.cells[-1].text
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip() == "DEA Classification"
    )
    assert "Yes\t\tNo" in dea_cell  # template checkbox text survives
    assert "->  Not scheduled [analyst: Second Analyst]" in dea_cell
    assert "Analyst input required" not in dea_cell


def test_overlay_note_on_priority_block_member(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A verified_absent block member keeps its 'No' marker; the overlay is an
    appended note inside the merged Priority Status cell."""
    template = tmp_path / "real-shaped.docx"
    _real_shaped_template(template)
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result, template_path=template, inputs=_inputs(drug_shortage="re-check next quarter")
    )
    doc = Document(BytesIO(data))
    block_cell = next(
        r.cells[-1].text
        for t in doc.tables
        for r in t.rows
        if r.cells and r.cells[0].text.strip() == "Priority Status"
    )
    assert "Drug Shortages: On Shortage List? Y/N  ->  No" in block_cell
    assert "Analyst note (Second Analyst): re-check next quarter" in block_cell


def test_overlay_analyst_inputs_appendix(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result,
        template_path=tmp_path / "missing.docx",
        inputs=_inputs(rd_center="Site A", product_name="verify with brand team"),
    )
    doc = Document(BytesIO(data))
    assert "Analyst inputs" in {p.text for p in doc.paragraphs}
    appendix = next(
        t
        for t in doc.tables
        if t.rows and [c.text for c in t.rows[0].cells] == ["Cell", "Value", "Author", "Updated at"]
    )
    rows = [[c.text for c in r.cells] for r in appendix.rows[1:]]
    assert ["R&D Center", "Site A", "Second Analyst", UPDATED_AT] in rows
    assert ["Product Name", "verify with brand team", "Second Analyst", UPDATED_AT] in rows
    assert len(rows) == 2


def test_overlay_fallback_path_parity(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """The from-scratch fallback applies the identical overlay treatment."""
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result,
        template_path=tmp_path / "missing.docx",
        inputs=_inputs(rd_center="Site A", product_name="verify with brand team"),
    )
    doc = Document(BytesIO(data))
    assert FALLBACK_MARKER in {p.text for p in doc.paragraphs}  # still the loud fallback
    section_rows = {
        r.cells[0].text: r.cells[1].text
        for t in doc.tables
        if t.rows and [c.text for c in t.rows[0].cells][:3] == ["Cell", "Value", "Status"]
        for r in t.rows[1:]
    }
    assert section_rows["R&D Center"] == "Site A [analyst: Second Analyst]"
    value_cell = section_rows["Product Name"]
    assert "ALBUTEROL SULFATE" in value_cell
    assert "Analyst note (Second Analyst): verify with brand team" in value_cell


def test_no_inputs_rendering_unchanged(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """inputs=None must render exactly as before the overlay existed: no analyst
    strings anywhere, and byte-identical document XML to an empty overlay."""
    template = tmp_path / "template.docx"
    _synthetic_template(template)
    result = _build(monkeypatch)
    plain = write_whitepaper_docx(result, template_path=template)
    for text in (_all_text(plain), _all_text(write_whitepaper_docx(result, template_path=None))):
        assert "[analyst:" not in text
        assert "Analyst note" not in text
        assert "Analyst inputs" not in text
    empty = write_whitepaper_docx(result, template_path=template, inputs={})
    assert _document_xml(plain) == _document_xml(empty)


def test_blank_overlay_value_never_renders_as_analyst_text(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A blank/whitespace overlay entry is ignored on every surface (INV-5)."""
    result = _build(monkeypatch)
    data = write_whitepaper_docx(
        result, template_path=tmp_path / "missing.docx", inputs=_inputs(rd_center="   ")
    )
    text = _all_text(data)
    assert "[analyst:" not in text
    assert "Analyst note" not in text
    assert "Analyst inputs" not in text
