"""Render a populated white paper as a Word document.

Two paths, per the feature contract:

1. **Real-template fill** — open ``settings.whitepaper_template_path`` with
   python-docx, locate each label cell in its tables and write the value into
   the adjacent value cell. Checkbox-style lines get an appended explicit
   ``-> Yes`` / ``-> No`` / ``-> Analyst input required`` marker (we do not try
   to tick the template's symbol checkboxes).
2. **From-scratch fallback** — when the template file is absent (CI), build a
   structurally-equivalent document from the registry. The fallback is never
   silent: it logs the missed path and stamps a visible marker line under the
   heading (``FALLBACK_MARKER``).

Both paths append a Provenance appendix table (cell -> source / locator /
fetched_at). The docx tests use a synthetic in-test template fixture, so CI
passes without the gitignored Word file.

Analyst overlay (Phase 2): ``write_whitepaper_docx`` takes an optional
``inputs`` map (cell_id -> {value, author, updated_at}) from the durable run
store. The generated layer renders EXACTLY as without inputs (INV-3): an
analyst value only ever FILLS a cell whose status is analyst_input_required
(visibly attributed), and on a populated / verified_absent cell it is APPENDED
as a distinct "Analyst note" paragraph -- a human note never replaces or
restyles a cited value. An "Analyst inputs" appendix follows the provenance
appendix whenever inputs exist.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell

from regwatch.common.logging import get_logger
from regwatch.whitepaper.template import (
    CHECKBOX_CELL_IDS,
    CellSpec,
    section_order,
    specs_for_section,
)

log = get_logger(__name__)

# Rendered under the heading whenever the document is built WITHOUT the real
# CRA template, so a fallback render can never pass for the official form.
FALLBACK_MARKER = "(generated without the official CRA template file)"

# Aliases mapping a registry cell id -> the normalized template label text it
# appears under. Matching is by PREFIX (longest alias first), because the real
# template suffixes labels with parentheticals ("Salable Unit (from Sales &
# Marketing)"). The template merges "Dosage Form; Route" into one row, so the
# dosage_form cell claims it and the route value is appended to the same cell.
_LABEL_ALIASES: dict[str, tuple[str, ...]] = {
    "product_name": ("product name",),
    "dosage_form": ("dosage form route", "dosage form"),
    "route": ("dosage form route", "route"),
    "strengths": ("strengths",),
    "rd_center": ("rd center", "r d center"),
    "priority_status": ("priority status",),
    "patents": ("patents ob review", "patents"),
    "first_to_market": ("if pi eligible for first to market",),
    "eftf": ("if piv eligible for eftf",),
    "drug_shortage": ("on drug shortage list", "drug shortages"),
    "combination_product": ("combination product",),
    "rld": ("reference listed drug rld", "reference listed drug"),
    "rs": ("reference standard rs", "reference standard"),
    "proprietary_name": ("proprietary name",),
    "rld_strength": ("rld strength",),
    "nda_number": ("nda", "nda number"),
    "nda_holder": ("nda holder",),
    "indication": ("indication",),
    "rems": ("rems",),
    "restricted_distribution": ("restricted distribution",),
    "labeling_images": ("labeling images",),
    "epc": ("established pharmacologic class",),
    "plr_format": ("plr physician labeling rule format", "plr format"),
    "usp_monograph": ("usp monograph",),
    "pllr_format": ("pllr pregnancy and lactation labeling rule format", "pllr format"),
    "pregnancy_registry": ("pregnancy registry contact detail",),
    "salable_unit": ("salable unit",),
    "packaging": ("packaging configurations", "packaging configuration s"),
    "labeling_carveouts": ("labeling carveouts",),
    "emergency_use": ("emergency use",),
    "dea_classification": ("dea classification",),
    "be_guidance_available": ("be guidance available",),
    "requirements": ("requirements",),
    "proposed_strategy": ("proposed strategy",),
    "in_vivo_be_studies": ("in vivo be studies",),
    "q1_q2_assessment": (
        "qualitatively q1 and quantitatively q2 assessment",
        "q1 q2 assessment",
    ),
    "tablet_scoring": ("tablet scoring",),
    "threshold_analysis": ("threshold analysis",),
    "human_factor_studies": ("human factor studies",),
    "dissolution_testing": ("dissolution testing",),
    "biocompatibility_studies": ("biocompatibility studies",),
    "toxicology_studies": ("toxicology studies",),
    "prepared_by": ("prepared by",),
    "reviewed_by": ("reviewed by",),
    "labeling_approved_by": ("labeling approved by",),
    "approved_by": ("approved by",),
}

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def docx_media_type() -> str:
    return _DOCX_MIME


def _norm(text: str) -> str:
    """Lowercase, drop punctuation, collapse whitespace — for label matching."""
    cleaned = []
    for ch in text.lower():
        cleaned.append(ch if ch.isalnum() or ch.isspace() else " ")
    return " ".join("".join(cleaned).split())


def _aliases_longest_first() -> list[tuple[str, str]]:
    """(alias, cell_id) pairs, longest alias first; first registration wins a tie."""
    pairs: list[tuple[str, str]] = []
    seen: set[str] = set()
    for cell_id, aliases in _LABEL_ALIASES.items():
        for alias in aliases:
            if alias not in seen:
                seen.add(alias)
                pairs.append((alias, cell_id))
    pairs.sort(key=lambda pair: len(pair[0]), reverse=True)
    return pairs


def _match_cell_id(label: str, aliases_longest_first: list[tuple[str, str]]) -> str | None:
    """Resolve a normalized label by exact or word-prefix match.

    The real template suffixes several labels ("Pregnancy Registry Contact
    Detail (if required)"), so exact equality would miss real rows; longest
    alias first keeps "nda holder" from being swallowed by "nda".
    """
    for alias, cell_id in aliases_longest_first:
        if label == alias or label.startswith(alias + " "):
            return cell_id
    return None


# The real template merges the Patents / PI-PIV / First-to-Market / eFTF /
# Drug-Shortages checkbox block into the single "Priority Status" value cell;
# these cells are appended there as markers, never overwriting the block.
_PRIORITY_BLOCK_MEMBER_IDS: tuple[str, ...] = (
    "priority_status",
    "patents",
    "first_to_market",
    "eftf",
    "drug_shortage",
)


def _cell_value(cell: dict[str, Any]) -> str | None:
    """The cell's renderable value — empty/whitespace reads as no value (C3).

    The populator's choke point already rejects empty values; this is the
    writer-side defense so a blank can never render as a populated cell.
    """
    value = cell.get("value")
    if isinstance(value, str) and value.strip():
        return value
    return None


# The analyst overlay map handed down from the run store:
# cell_id -> {"value": str, "author": str, "updated_at": str}.
Inputs = dict[str, dict[str, Any]] | None


def _overlay(inputs: Inputs, cell_id: str) -> tuple[str, str] | None:
    """(value, author) when the overlay carries a usable value for this cell.

    Defensive at the render boundary: the store never persists a blank value,
    but a blank must still never render as analyst text (INV-5 mirror of
    ``_cell_value``).
    """
    if not inputs:
        return None
    entry = inputs.get(cell_id)
    if not isinstance(entry, dict):
        return None
    value = entry.get("value")
    if not isinstance(value, str) or not value.strip():
        return None
    author = entry.get("author")
    author_name = author if isinstance(author, str) and author.strip() else "unknown"
    return value, author_name


def _analyst_fill(cell: dict[str, Any], inputs: Inputs) -> str | None:
    """The attributed analyst text, ONLY for a manual cell awaiting input.

    Keyed on status so an overlay can never substitute for a generated value:
    populated / verified_absent cells keep their cited rendering and take the
    note via ``_analyst_note`` instead (INV-3).
    """
    if cell["status"] != "analyst_input_required":
        return None
    overlay = _overlay(inputs, cell["id"])
    if overlay is None:
        return None
    value, author = overlay
    return f"{value} [analyst: {author}]"


def _analyst_note(cell: dict[str, Any], inputs: Inputs) -> str | None:
    """The appended note paragraph for a cell that already carries a
    generated rendering -- never a replacement (INV-3)."""
    if cell["status"] == "analyst_input_required":
        return None
    overlay = _overlay(inputs, cell["id"])
    if overlay is None:
        return None
    value, author = overlay
    return f"Analyst note ({author}): {value}"


def _marker(cell: dict[str, Any], inputs: Inputs = None) -> str:
    status = cell["status"]
    if status == "verified_absent":
        return "No"
    if status == "analyst_input_required":
        return _analyst_fill(cell, inputs) or "Analyst input required"
    return _cell_value(cell) or "Yes"


def _render_value(cell: dict[str, Any], inputs: Inputs = None) -> str:
    value = _cell_value(cell)
    if value is not None:
        return value
    if cell["status"] == "verified_absent":
        return "No (verified absent)"
    return _analyst_fill(cell, inputs) or "Analyst input required"


def _cells_by_id(result: dict[str, Any]) -> dict[str, dict[str, Any]]:
    out: dict[str, dict[str, Any]] = {}
    for section in result["sections"]:
        for cell in section["cells"]:
            out[cell["id"]] = cell
    return out


def write_whitepaper_docx(
    result: dict[str, Any], *, template_path: Path | None, inputs: Inputs = None
) -> bytes:
    """Render ``result`` to .docx bytes (real template if present, else scratch).

    ``inputs`` is the attributed analyst overlay (cell_id -> {value, author,
    updated_at}); ``None`` renders exactly as before the overlay existed.
    """
    if template_path is not None and template_path.exists():
        doc = _fill_template(template_path, result, inputs)
    else:
        # Never silent: name the path that was missed and fall back loudly —
        # the fallback document itself carries a visible marker line.
        log.warning(
            "whitepaper_template_missing",
            template_path=str(template_path) if template_path is not None else "(not configured)",
        )
        doc = _build_from_scratch(result, inputs)
    _append_provenance(doc, result)
    _append_analyst_inputs(doc, result, inputs)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _fill_template(template_path: Path, result: dict[str, Any], inputs: Inputs) -> DocxDocument:
    doc = Document(str(template_path))
    cells = _cells_by_id(result)
    aliases = _aliases_longest_first()
    filled: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            row_cells = row.cells
            if not row_cells:
                continue
            label = _norm(row_cells[0].text)
            cell_id = _match_cell_id(label, aliases)
            # Fill each cell once: the real template repeats some label rows
            # (Combination Product appears twice), and a second write would
            # land the marker in the wrong cell.
            if cell_id is None or cell_id not in cells or cell_id in filled:
                continue
            target = row_cells[-1]
            if cell_id == "priority_status":
                _fill_priority_block(target, cells, filled, inputs)
                continue
            _write_value_cell(
                target, cells[cell_id], same_as_label=target is row_cells[0], inputs=inputs
            )
            filled.add(cell_id)
            if cell_id == "dosage_form" and label.startswith("dosage form route"):
                route = cells.get("route")
                if route is not None and "route" not in filled:
                    target.add_paragraph(f"Route: {_render_value(route, inputs)}")
                    _append_note_paragraph(target, route, inputs)
                    filled.add("route")
    # Cells with no matching template row are recorded in the provenance
    # appendix and an "unmapped values" table so nothing is silently dropped.
    _append_unmapped(doc, result, filled, inputs)
    return doc


def _append_note_paragraph(target: _Cell, cell: dict[str, Any], inputs: Inputs) -> None:
    """Append the analyst note under a generated rendering, when one exists."""
    note = _analyst_note(cell, inputs)
    if note is not None:
        target.add_paragraph(note)


def _fill_priority_block(
    target: _Cell, cells: dict[str, dict[str, Any]], filled: set[str], inputs: Inputs
) -> None:
    """Append per-cell markers into the merged Priority Status checkbox block.

    The block cell carries the template's own Patents / PI-PIV / eFTF /
    Drug-Shortages checkbox lines; overwriting it would delete them, so every
    member rides in as an appended "label -> marker" paragraph instead. An
    analyst overlay on a member follows the same discipline: the manual cells
    take the attributed text as their marker, the generated ones keep their
    marker and gain an appended note (INV-3).
    """
    for member in _PRIORITY_BLOCK_MEMBER_IDS:
        cell = cells.get(member)
        if cell is None or member in filled:
            continue
        target.add_paragraph(f"{cell['label']}  ->  {_marker(cell, inputs)}")
        _append_note_paragraph(target, cell, inputs)
        filled.add(member)


def _write_value_cell(
    target: _Cell, cell: dict[str, Any], *, same_as_label: bool, inputs: Inputs = None
) -> None:
    is_checkbox = cell["id"] in CHECKBOX_CELL_IDS
    if same_as_label:
        # Merged single-cell row — append rather than wipe the label/checkboxes.
        target.add_paragraph(f"{cell['label']}: {_render_or_marker(cell, is_checkbox, inputs)}")
        _append_note_paragraph(target, cell, inputs)
        return
    if is_checkbox:
        existing = target.text.strip()
        suffix = f"  ->  {_marker(cell, inputs)}"
        if existing:
            target.add_paragraph(suffix.strip())
        else:
            target.text = suffix.strip()
        _append_note_paragraph(target, cell, inputs)
        return
    target.text = _render_value(cell, inputs)
    _append_note_paragraph(target, cell, inputs)


def _render_or_marker(cell: dict[str, Any], is_checkbox: bool, inputs: Inputs = None) -> str:
    return _marker(cell, inputs) if is_checkbox else _render_value(cell, inputs)


def _append_unmapped(
    doc: DocxDocument, result: dict[str, Any], filled: set[str], inputs: Inputs
) -> None:
    missing = [
        cell
        for section in result["sections"]
        for cell in section["cells"]
        if cell["id"] not in filled
    ]
    if not missing:
        return
    doc.add_page_break()
    doc.add_heading("Additional populated values (not mapped to a template cell)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Cell"
    hdr[1].text = "Value"
    hdr[2].text = "Status"
    for cell in missing:
        cells = table.add_row().cells
        cells[0].text = cell["label"]
        cells[1].text = _render_value(cell, inputs)
        _append_note_paragraph(cells[1], cell, inputs)
        cells[2].text = cell["status"]


def _build_from_scratch(result: dict[str, Any], inputs: Inputs = None) -> DocxDocument:
    doc = Document()
    spine = result.get("spine", {})
    doc.add_heading("CRA White Paper", level=0)
    doc.add_paragraph(FALLBACK_MARKER)
    doc.add_paragraph(
        f"{spine.get('application_type', '')} {spine.get('application_number', '')} — "
        f"{spine.get('ingredient', '')}"
    )
    if spine.get("approved_label_document_id"):
        doc.add_paragraph(f"Drugs@FDA label: {spine['approved_label_document_id']}")
    cells = _cells_by_id(result)
    for title in section_order():
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Cell"
        hdr[1].text = "Value"
        hdr[2].text = "Status"
        for spec in specs_for_section(title):
            cell = cells.get(spec.id)
            if cell is None:  # pragma: no cover - registry/result drift guard
                continue
            row = table.add_row().cells
            row[0].text = _cell_label(spec, cell)
            row[1].text = _render_value(cell, inputs)
            _append_note_paragraph(row[1], cell, inputs)
            row[2].text = cell["status"]
    return doc


def _cell_label(spec: CellSpec, cell: dict[str, Any]) -> str:
    suffix = " (Yes/No)" if spec.id in CHECKBOX_CELL_IDS else ""
    return f"{cell['label']}{suffix}"


def _append_provenance(doc: DocxDocument, result: dict[str, Any]) -> None:
    doc.add_page_break()
    doc.add_heading("Provenance appendix", level=1)
    doc.add_paragraph(
        "Every populated value below is traceable to a fetched FDA source row. "
        "Cells marked 'analyst input required' carry evidence but no generated value."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Cell"
    hdr[1].text = "Source"
    hdr[2].text = "Locator"
    hdr[3].text = "Fetched at"
    for section in result["sections"]:
        for cell in section["cells"]:
            if not cell["evidence"]:
                continue
            for ev in cell["evidence"]:
                row = table.add_row().cells
                row[0].text = cell["label"]
                row[1].text = str(ev.get("source") or "")
                row[2].text = str(ev.get("locator") or "")
                row[3].text = str(ev.get("fetched_at") or "")


def _append_analyst_inputs(doc: DocxDocument, result: dict[str, Any], inputs: Inputs) -> None:
    """Second appendix table: the attributed analyst overlay, when one exists.

    Registry order (via the result's sections) so the table reads like the
    document; any overlay key not in the result (drift guard) still lands at
    the end rather than being silently dropped (INV-4).
    """
    if not inputs:
        return
    ordered: list[tuple[str, str]] = []  # (cell_id, label)
    seen: set[str] = set()
    for section in result["sections"]:
        for cell in section["cells"]:
            if cell["id"] in inputs and cell["id"] not in seen:
                ordered.append((cell["id"], cell["label"]))
                seen.add(cell["id"])
    for cell_id in sorted(set(inputs) - seen):
        ordered.append((cell_id, cell_id))
    rows: list[tuple[str, str, str, str]] = []
    for cell_id, label in ordered:
        overlay = _overlay(inputs, cell_id)
        if overlay is None:  # blank/malformed entry: never render as analyst text
            continue
        value, author = overlay
        entry = inputs.get(cell_id) or {}
        rows.append((label, value, author, str(entry.get("updated_at") or "")))
    if not rows:
        return
    doc.add_page_break()
    doc.add_heading("Analyst inputs", level=1)
    doc.add_paragraph(
        "Attributed human input recorded against this run. These values are "
        "analyst-entered, never generated; populated cells above keep their "
        "cited values and carry these only as notes."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Cell"
    hdr[1].text = "Value"
    hdr[2].text = "Author"
    hdr[3].text = "Updated at"
    for label, value, author, updated_at in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        row[2].text = author
        row[3].text = updated_at
