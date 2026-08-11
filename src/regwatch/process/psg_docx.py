"""Renders a rebuilt PSG as a Word document.

The studio's working documents are .docx files, and an analyst who opens a
reference PSG beside one expects to be able to take it away in the same
format -- into a submission working folder, a review pack, an email. This
module turns the blocks that `process.psg_document` rebuilt into those bytes.

Nothing is stored. The document is generated per request from text already in
the database, so there is no file to invalidate when FDA revises a PSG and no
second copy of the corpus to keep in step.

The provenance paragraph is not decoration. This file is a machine extraction
of a published PDF: the words are FDA's, the layout, pagination and tables are
not, so the document says where it came from and which artifact is
authoritative. A Word file that looked like an FDA original without saying so
would be the one genuinely dangerous thing this feature could produce.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.shared import Pt

from regwatch.process.psg_document import PsgBlock

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Word maps these to its own built-in styles; a heading level rather than bold
# body text is what makes the file navigable in Word's pane.
_TITLE_SIZE_PT = 16
_META_SIZE_PT = 9


@dataclass(frozen=True)
class PsgDocxMeta:
    """The identifying fields printed under the title and in provenance."""

    active_ingredient: str
    dosage_form: str | None
    route: str | None
    appl_no: str | None
    psg_type: str
    recommended_date: str | None
    source_url: str


def safe_file_stem(name: str) -> str:
    """A file-name stem safe for a Content-Disposition header.

    Everything outside ``[A-Za-z0-9._-]`` collapses to an underscore, which
    rules out the quote and CR/LF characters a header value must never carry
    (the same discipline the PDF and whitepaper routes apply to an
    application number, applied here to a drug name that can hold commas,
    semicolons and parentheses).
    """
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")
    return cleaned or "psg"


def _describe(meta: PsgDocxMeta) -> str:
    """The one-line descriptor under the title: form, route, type, date."""
    parts = [p for p in (meta.dosage_form, meta.route) if p]
    parts.append("Final guidance" if meta.psg_type == "final" else "Draft guidance")
    if meta.recommended_date:
        parts.append(f"Recommended {meta.recommended_date}")
    if meta.appl_no:
        parts.append(f"PSG_{meta.appl_no}")
    return "  |  ".join(parts)


def write_psg_docx(
    blocks: list[PsgBlock],
    meta: PsgDocxMeta,
    *,
    truncated: bool = False,
) -> bytes:
    """Renders one PSG's blocks as .docx bytes.

    Args:
        blocks: The document body, in reading order, from
            `psg_document.build_body`.
        meta: The identifying fields for the title block and provenance note.
        truncated: True when the rebuild hit its runaway guard. Stated in the
            document rather than dropped, so a short file is never mistaken
            for a short guidance.

    Returns:
        The .docx file as bytes. Never empty: a PSG with no rebuilt blocks
        still produces the title, the descriptor and the provenance note, so
        the download explains itself instead of failing silently.
    """
    document = Document()

    heading = document.add_paragraph()
    run = heading.add_run(f"PSG: {meta.active_ingredient}")
    run.bold = True
    run.font.size = Pt(_TITLE_SIZE_PT)

    descriptor = document.add_paragraph()
    descriptor_run = descriptor.add_run(_describe(meta))
    descriptor_run.italic = True
    descriptor_run.font.size = Pt(_META_SIZE_PT)

    for block in blocks:
        if block.type == "title":
            document.add_heading(block.text, level=1)
        elif block.type == "h2":
            document.add_heading(block.text, level=2)
        elif block.type == "meta":
            paragraph = document.add_paragraph()
            meta_run = paragraph.add_run(block.text)
            meta_run.font.size = Pt(_META_SIZE_PT)
        else:
            document.add_paragraph(block.text)

    if truncated:
        note = document.add_paragraph()
        note_run = note.add_run(
            "This extract was truncated: the source document produced more "
            "sections than this renderer emits. Use the FDA PDF for the "
            "complete guidance."
        )
        note_run.bold = True

    document.add_paragraph()
    provenance = document.add_paragraph()
    provenance_run = provenance.add_run(
        "Text extracted from the FDA product-specific guidance PDF at "
        f"{meta.source_url} . The PDF published by FDA is the authoritative "
        "document; this file reproduces its text for review and is not an "
        "FDA-issued document. Layout, pagination and any tables in the "
        "original are not reproduced."
    )
    provenance_run.italic = True
    provenance_run.font.size = Pt(_META_SIZE_PT)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
